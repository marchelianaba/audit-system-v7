#!/usr/bin/env python3
"""
render_lhp.py — Renderer LHP DOCX placeholder-driven (v4.0.4).

Konsumsi `_KKP/temuan.json` + `context.md` + `_PKP/sasaran-assignment.json`
+ template `templates/_skeleton-lhp/template-lhp-[jenis].docx`, lalu lakukan
find-replace placeholder `{{...}}` untuk menghasilkan
`_LHP/LHP-SUBSTANSI-[nomor-st-slug].docx`.

Placeholder yang didukung:
  Identitas: NOMOR_ST, TANGGAL_ST, OBYEK, PENERIMA_LHP, PERIODE_PELAKSANAAN,
             BULAN_TAHUN, NOMOR_LHR, NOMOR_NOTA_DINAS, TANGGAL_NOTA_DINAS,
             HAL_LHR, JUDUL_LHR_LINE_1..3, JUDUL_LHR_INLINE, DASAR_PERMINTAAN,
             NAMA_AUDITI, TANGGAL_EXIT_MEETING
  Tim     : NAMA_INSPEKTUR, NIP_INSPEKTUR, TTD_INSPEKTUR, TEMBUSAN_LIST
  Konten  : DASAR_HUKUM_LIST, TUJUAN_REVIU, SASARAN_LIST, RUANG_LINGKUP,
             METODOLOGI_REVIU, GAMBARAN_UMUM, HASIL_REVIU_INTRO,
             HASIL_REVIU_LOOP, SIMPULAN_REVIU

Yang tidak diisi otomatis (harus dilengkapi auditor manual via INTEGRAL/SIMWAS):
  NOMOR_LHR, NOMOR_NOTA_DINAS, TANGGAL_NOTA_DINAS, TTD_INSPEKTUR

Contoh:
  python3 scripts/render_lhp.py --penugasan penugasan/[nama] \
      --rekomendasi-file penugasan/[nama]/_LHP/rekomendasi.json \
      --judul "Pengadaan DC/DRC PSrE Induk Tahun 2026" \
      --auditi "Direktorat Pengawasan Sertifikasi dan Transaksi Elektronik"
"""
from __future__ import annotations

import argparse
import copy
import json
import re
import sys
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.stderr.write("python-docx tidak terpasang. Jalankan: pip install python-docx\n")
    sys.exit(2)


JENIS_LABEL = {
    "reviu-pengadaan": "Reviu Pengadaan",
    "reviu-rka-kl": "Reviu Rencana Kerja dan Anggaran",
    "audit-pengadaan": "Audit Pengadaan",
    "audit-kinerja": "Audit Kinerja",
    "evaluasi-sakip": "Evaluasi SAKIP",
    "evaluasi-spip": "Evaluasi SPIP",
    "evaluasi-reformasi-birokrasi": "Evaluasi Reformasi Birokrasi",
    "evaluasi-manajemen-risiko": "Evaluasi Manajemen Risiko",
    "pemantauan-pengadaan": "Pemantauan Pengadaan",
    "pemantauan-tindak-lanjut": "Pemantauan Tindak Lanjut Hasil Pengawasan",
}


# ---------- Context parser ----------
def parse_context(context_path: Path) -> dict:
    """Ekstrak field dari tabel context.md."""
    out = {}
    if not context_path.exists():
        return out
    text = context_path.read_text(encoding="utf-8")
    field_map = {
        "nomor st": "nomor_st", "tanggal st": "tanggal_st",
        "objek": "obyek", "obyek": "obyek",
        "paket pengadaan": "paket",
        "tahun anggaran": "tahun_anggaran",
        "periode pelaksanaan": "periode",
        "jangka waktu": "jangka_waktu",
        "tingkat risiko": "tingkat_risiko",
        "tingkat keyakinan": "tingkat_keyakinan",
        "penerima lhp": "penerima_lhp",
        "dasar penugasan": "dasar_penugasan",
    }
    for line in text.splitlines():
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|") if c.strip()]
        if len(cells) >= 2:
            k_norm = cells[0].lower()
            for k_pattern, target_field in field_map.items():
                if k_pattern in k_norm:
                    if target_field not in out or not out[target_field]:
                        out[target_field] = cells[1]
                    break
    # Tujuan + Ruang Lingkup (inline format Tujuan: ... / Ruang Lingkup: ...)
    m = re.search(r"^Tujuan\s*:\s*(.+?)(?=\n\n|\n##|\nRuang Lingkup)", text, re.MULTILINE | re.DOTALL)
    if m:
        out["tujuan"] = m.group(1).strip()
    m = re.search(r"^Ruang Lingkup\s*:\s*(.+?)(?=\n\n|\n##)", text, re.MULTILINE | re.DOTALL)
    if m:
        out["ruang_lingkup"] = m.group(1).strip()
    # Ringkasan Obyek — heading ## Ringkasan Obyek diikuti paragraf isi
    m = re.search(r"##\s*Ringkasan Obyek\s*\n+([\s\S]+?)(?=\n##|\Z)", text, re.IGNORECASE)
    if m:
        out["ringkasan_obyek"] = m.group(1).strip()
    # Tim — tabel
    tim = []
    in_tim = False
    for line in text.splitlines():
        if line.strip().startswith("## Tim"):
            in_tim = True
            continue
        if in_tim and line.startswith("|"):
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if len(cells) >= 5 and cells[0].isdigit():
                tim.append({"no": cells[0], "nama": cells[1], "nip": cells[2],
                           "jabatan": cells[3], "jabfung": cells[4]})
        elif in_tim and not line.startswith("|") and line.strip() and not line.startswith("---"):
            in_tim = False
    out["tim"] = tim
    return out


# ---------- Find-replace ----------
def replace_in_paragraph(p, mapping: dict):
    """Replace {{KEY}} placeholders in a paragraph, preserving formatting of
    the first run that contains the placeholder."""
    full = "".join(r.text for r in p.runs)
    if "{{" not in full:
        return False
    new = full
    for k, v in mapping.items():
        new = new.replace("{{" + k + "}}", str(v))
    if new == full:
        return False
    # Replace all runs with single run carrying full new text, preserving
    # formatting of first run.
    if p.runs:
        first = p.runs[0]
        first.text = new
        for r in p.runs[1:]:
            r.text = ""
    return True


def replace_in_doc(doc, mapping: dict):
    for p in doc.paragraphs:
        replace_in_paragraph(p, mapping)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, mapping)


# ---------- Loop helpers ----------
def expand_paragraph_placeholder_to_blocks(doc, placeholder: str, blocks: list):
    """Replace SETIAP paragraph yang mengandung `placeholder` dengan list of
    paragraph blocks. Mengembalikan jumlah occurrences yang berhasil di-expand.
    """
    targets = [p for p in doc.paragraphs if placeholder in p.text]
    if not targets:
        return 0

    def make_p(text, fmt=None):
        fmt = fmt or {}
        new = doc.add_paragraph()
        if fmt.get("align") == "justify":
            new.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif fmt.get("align") == "center":
            new.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif fmt.get("align") == "right":
            new.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if fmt.get("indent"):
            new.paragraph_format.left_indent = Cm(fmt["indent"])
        r = new.add_run(text)
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
        r.font.size = Pt(fmt.get("size", 11))
        r.font.name = "Arial"
        return new

    expanded = 0
    for target_p in targets:
        target_el = target_p._element
        new_paras = [make_p(t, f) for t, f in blocks]
        for new_p in new_paras:
            target_el.addprevious(new_p._element)
        target_el.getparent().remove(target_el)
        expanded += 1
    return expanded


# ---------- Main composition ----------
def build_mapping(pen_dir: Path, args) -> tuple[dict, list, dict, dict]:
    """Bangun mapping placeholder dari context + temuan + sasaran."""
    ctx = parse_context(pen_dir / "context.md")
    temuan_path = pen_dir / "_KKP" / "temuan.json"
    sasaran_path = pen_dir / "_PKP" / "sasaran-assignment.json"
    if not temuan_path.exists():
        sys.stderr.write(f"temuan.json tidak ditemukan: {temuan_path}\n"); sys.exit(1)
    if not sasaran_path.exists():
        sys.stderr.write(f"sasaran-assignment.json tidak ditemukan: {sasaran_path}\n"); sys.exit(1)
    kkp = json.loads(temuan_path.read_text(encoding="utf-8"))
    sa = json.loads(sasaran_path.read_text(encoding="utf-8"))

    # Rekomendasi mapping (per id_temuan)
    rek = {}
    if args.rekomendasi_file and Path(args.rekomendasi_file).exists():
        rek = json.loads(Path(args.rekomendasi_file).read_text(encoding="utf-8"))

    jenis = kkp["penugasan"]["jenis_pengawasan"]
    obyek = ctx.get("obyek") or kkp["penugasan"]["obyek"]
    penerima_lhp = ctx.get("penerima_lhp") or args.penerima or "[DIISI AUDITOR]"
    nomor_st = ctx.get("nomor_st") or kkp["penugasan"]["nomor_st"]
    tanggal_st = ctx.get("tanggal_st") or kkp["penugasan"]["tanggal_st"]
    periode = ctx.get("periode") or "[DIISI AUDITOR]"

    # Inspektur (PM)
    pm = next((m for m in ctx.get("tim", []) if "Pengendali Mutu" in m.get("jabatan", "")), None)
    nama_inspektur = pm["nama"] if pm else "[DIISI AUDITOR]"
    nip_inspektur = pm["nip"] if pm else "[DIISI]"

    judul_lhr = args.judul or f"{kkp['penugasan']['obyek']}"
    # Pecah judul jadi 3 baris kalau panjang
    words = judul_lhr.split()
    if len(words) > 12:
        third = len(words) // 3
        line1 = " ".join(words[:third])
        line2 = " ".join(words[third:2*third])
        line3 = " ".join(words[2*third:])
    else:
        line1 = judul_lhr
        line2 = ""
        line3 = ""

    nama_auditi = args.auditi or obyek

    # Bulan/Tahun untuk halaman judul
    bulan_id = ["JANUARI","FEBRUARI","MARET","APRIL","MEI","JUNI","JULI","AGUSTUS","SEPTEMBER","OKTOBER","NOVEMBER","DESEMBER"]
    today = datetime.now()
    bulan_tahun = f"{bulan_id[today.month-1]} {today.year}"

    mapping = {
        "NOMOR_ST": nomor_st,
        "TANGGAL_ST": tanggal_st,
        "OBYEK": obyek,
        "PENERIMA_LHP": penerima_lhp,
        "PERIODE_PELAKSANAAN": periode,
        "BULAN_TAHUN": bulan_tahun,
        "NOMOR_LHR": "[DIISI AUDITOR — dari SIMWAS]",
        "NOMOR_NOTA_DINAS": "[DIISI AUDITOR]",
        "TANGGAL_NOTA_DINAS": "[DIISI AUDITOR]",
        "HAL_LHR": judul_lhr,
        "JUDUL_LHR_LINE_1": line1.upper() if line1 else "",
        "JUDUL_LHR_LINE_2": line2.upper() if line2 else "",
        "JUDUL_LHR_LINE_3": line3.upper() if line3 else "",
        "JUDUL_LHR_INLINE": judul_lhr,
        "DASAR_PERMINTAAN": args.dasar_permintaan or "Surat Tugas Inspektur Jenderal Nomor " + nomor_st,
        "NAMA_AUDITI": nama_auditi,
        "TANGGAL_EXIT_MEETING": args.tanggal_exit_meeting or "[DIISI AUDITOR]",
        "NAMA_INSPEKTUR": nama_inspektur,
        "NIP_INSPEKTUR": nip_inspektur,
        "TTD_INSPEKTUR": "[DIISI AUDITOR — TTD]",
        "TUJUAN_REVIU": ctx.get("tujuan", "[DIISI dari context.md atau argumen]"),
        "RUANG_LINGKUP": ctx.get("ruang_lingkup", "[DIISI dari context.md]"),
        "METODOLOGI_REVIU": (
            "Reviu dilaksanakan dengan melakukan penelaahan dokumen (desk review) "
            "atas seluruh data dukung yang diterima dari auditi, dipadukan dengan "
            "klarifikasi tertulis kepada PPK dan tim teknis. Tim juga memanfaatkan "
            "pipeline pre-digest dan cross-check otomatis audit-system-v4 untuk "
            "mendeteksi anomali struktural antar dokumen."
        ),
        "GAMBARAN_UMUM": args.gambaran_umum or ctx.get("ringkasan_obyek") or "[DIISI — gambaran obyek pengadaan, nilai HPS, mekanisme pengadaan]",
        "HASIL_REVIU_INTRO": (
            f"Berdasarkan penelaahan atas dokumen perencanaan, tim Inspektorat II "
            f"memperoleh {len(kkp['temuan'])} ({_terbilang(len(kkp['temuan']))}) "
            f"catatan reviu yang dikelompokkan ke dalam {len(sa['sasaran'])} "
            f"({_terbilang(len(sa['sasaran']))}) aspek sesuai sasaran pengawasan. "
            f"Catatan-catatan ini dirumuskan dengan paradigma reviu (Kondisi-Kriteria-"
            f"Akibat-Rekomendasi) dengan tingkat keyakinan terbatas sebagaimana "
            f"diatur dalam SAIPI dan Perlem LKPP 12/2021."
        ),
    }
    return mapping, kkp, sa, rek


def _terbilang(n: int) -> str:
    digits = ["nol","satu","dua","tiga","empat","lima","enam","tujuh","delapan",
              "sembilan","sepuluh","sebelas"]
    if 0 <= n < len(digits):
        return digits[n]
    return str(n)


def build_dasar_hukum_blocks(jenis: str) -> list:
    items = [
        "Peraturan Pemerintah Nomor 60 Tahun 2008 tentang Sistem Pengendalian Intern Pemerintah;",
    ]
    if "pengadaan" in jenis:
        items += [
            "Peraturan Presiden Nomor 16 Tahun 2018 sebagaimana diubah dengan Peraturan Presiden Nomor 12 Tahun 2021 tentang Pengadaan Barang/Jasa Pemerintah;",
            "Peraturan Lembaga Kebijakan Pengadaan Barang/Jasa Pemerintah Nomor 12 Tahun 2021 tentang Pedoman Pelaksanaan Pengadaan Barang/Jasa Pemerintah melalui Penyedia;",
        ]
    items += [
        "Standar Audit Intern Pemerintah Indonesia (SAIPI) AAIPI 2021 (PER-01/AAIPI/DPN/2021);",
        "Program Kerja Pengawasan Tahunan Inspektorat Jenderal Komdigi.",
    ]
    blocks = []
    for i, txt in enumerate(items, start=1):
        blocks.append((f"{i}. {txt}", {"align": "justify"}))
    return blocks


def build_sasaran_blocks(sa: dict) -> list:
    blocks = []
    for s in sa["sasaran"]:
        blocks.append((f"{s['sasaran_id']}. {s['deskripsi']}", {"align": "justify", "indent": 0.8}))
    return blocks


def build_hasil_reviu_blocks(kkp: dict, sa: dict, rekomendasi: dict, jenis: str) -> list:
    blocks = []
    per_sasaran = {}
    for t in kkp["temuan"]:
        per_sasaran.setdefault(t["sasaran_id"], []).append(t)

    aspek_letter_seq = ["F.1", "F.2", "F.3", "F.4", "F.5", "F.6", "F.7", "F.8"]

    for idx, sid in enumerate(sorted(per_sasaran.keys())):
        sasaran_obj = next(x for x in sa["sasaran"] if x["sasaran_id"] == sid)
        # Sub-heading aspek
        section_title = f"{aspek_letter_seq[idx]}. Aspek {sid} — {sasaran_obj['deskripsi']}"
        blocks.append((section_title, {"bold": True, "size": 12}))
        blocks.append((f"Sasaran: {sasaran_obj['deskripsi']}", {"italic": True, "indent": 0.5, "align": "justify"}))

        for ti, t in enumerate(per_sasaran[sid], start=1):
            blocks.append((f"{sid}.{ti} {t['judul_temuan']}", {"bold": True}))

            blocks.append(("Kondisi:", {"bold": True}))
            blocks.append((t["kondisi"], {"align": "justify", "indent": 0.3}))

            blocks.append(("Kriteria:", {"bold": True}))
            blocks.append((t["kriteria"], {"align": "justify", "indent": 0.3}))

            if jenis.startswith("audit") and t.get("sebab"):
                blocks.append(("Sebab:", {"bold": True}))
                blocks.append((t["sebab"], {"align": "justify", "indent": 0.3}))

            if "Akibat" in t.get("akibat", "") or t.get("akibat"):
                blocks.append(("Akibat:", {"bold": True}))
                blocks.append((t.get("akibat", "—"), {"align": "justify", "indent": 0.3}))

            blocks.append(("Rekomendasi:", {"bold": True}))
            rek = rekomendasi.get(t["id_temuan"], "[Rekomendasi disusun bersama Pengendali Teknis berdasarkan exit meeting]")
            blocks.append((rek, {"align": "justify", "indent": 0.3}))

            sumber = "; ".join(f"{ds['file']} hal. {ds['halaman']}" for ds in t.get("dokumen_sumber", []))
            blocks.append((f"Sumber dokumen: {sumber}", {"italic": True, "size": 9, "indent": 0.3, "align": "justify"}))
            blocks.append(("", {}))  # spacer
    return blocks


def build_simpulan(kkp: dict, sa: dict) -> str:
    n = len(kkp["temuan"])
    return (
        f"Berdasarkan hasil reviu terbatas yang kami lakukan, dokumen perencanaan "
        f"yang menjadi obyek reviu secara umum telah disusun untuk mendukung "
        f"pelaksanaan kegiatan auditi dan telah memuat persyaratan utama. Namun, "
        f"terdapat {n} ({_terbilang(n)}) catatan reviu yang perlu ditindaklanjuti "
        f"sebelum dimulainya tahap pemilihan penyedia/pelaksanaan kegiatan. "
        f"Dengan keterbatasan keyakinan reviu (limited assurance), kami tidak "
        f"memberikan opini final atas kelayakan harga maupun kelayakan teknis; "
        f"opini tersebut menjadi tanggung jawab Pejabat Pembuat Komitmen, Pokja "
        f"Pemilihan, dan tim teknis pengadaan pada saat tender dilaksanakan."
    )


def build_tembusan_blocks() -> list:
    items = [
        "Inspektur Jenderal Kementerian Komunikasi dan Digital;",
        "Sekretaris Inspektorat Jenderal;",
        "Arsip.",
    ]
    return [(f"{i}. {txt}", {"size": 10}) for i, txt in enumerate(items, start=1)]


def main() -> int:
    ap = argparse.ArgumentParser(description="Render LHP DOCX dari template + temuan.json")
    ap.add_argument("--penugasan", required=True)
    ap.add_argument("--rekomendasi-file", default=None,
                    help="Path JSON {id_temuan: 'rekomendasi text'}")
    ap.add_argument("--judul", default=None, help="Judul LHR (mis. 'Pengadaan DC/DRC PSrE Induk Tahun 2026')")
    ap.add_argument("--auditi", default=None, help="Nama auditi (default: dari obyek)")
    ap.add_argument("--penerima", default=None, help="Penerima LHP (override context)")
    ap.add_argument("--dasar-permintaan", default=None,
                    help="Dasar permintaan reviu (override default ST-only)")
    ap.add_argument("--gambaran-umum", default=None, help="Paragraf gambaran umum pengadaan")
    ap.add_argument("--tanggal-exit-meeting", default=None)
    ap.add_argument("--template", default=None,
                    help="Override template path (default: templates/_skeleton-lhp/template-lhp-[jenis].docx)")
    ap.add_argument("--out", default=None,
                    help="Output path (default: _LHP/LHP-SUBSTANSI-[nomor-st-slug].docx)")
    args = ap.parse_args()

    pen_dir = Path(args.penugasan)
    if not pen_dir.exists():
        sys.stderr.write(f"Folder tidak ada: {pen_dir}\n"); return 1

    mapping, kkp, sa, rek = build_mapping(pen_dir, args)
    jenis = kkp["penugasan"]["jenis_pengawasan"]

    # Resolve template
    if args.template:
        tpl_path = Path(args.template)
    else:
        # Cari skeleton template di dua kemungkinan path (relative ke cwd dan absolute)
        candidates = [
            Path("templates/_skeleton-lhp") / f"template-lhp-{jenis}.docx",
            Path(__file__).resolve().parent.parent / "templates/_skeleton-lhp" / f"template-lhp-{jenis}.docx",
        ]
        tpl_path = next((p for p in candidates if p.exists()), None)
    if tpl_path is None or not tpl_path.exists():
        sys.stderr.write(f"Template skeleton tidak ditemukan untuk jenis '{jenis}'\n")
        return 1

    # Resolve output
    if args.out:
        out_path = Path(args.out)
    else:
        slug = re.sub(r"[^A-Za-z0-9]+", "-", mapping["NOMOR_ST"]).strip("-")
        out_path = pen_dir / "_LHP" / f"LHP-SUBSTANSI-{slug}.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # Load template
    doc = Document(str(tpl_path))

    # 1. Replace simple placeholders
    replace_in_doc(doc, mapping)

    # 2. Expand list/loop placeholders
    expand_paragraph_placeholder_to_blocks(doc, "{{DASAR_HUKUM_LIST}}", build_dasar_hukum_blocks(jenis))
    expand_paragraph_placeholder_to_blocks(doc, "{{SASARAN_LIST}}", build_sasaran_blocks(sa))
    expand_paragraph_placeholder_to_blocks(doc, "{{HASIL_REVIU_LOOP}}", build_hasil_reviu_blocks(kkp, sa, rek, jenis))
    # Simpulan as paragraph (single text)
    replace_in_doc(doc, {"SIMPULAN_REVIU": build_simpulan(kkp, sa)})
    # Tembusan
    expand_paragraph_placeholder_to_blocks(doc, "{{TEMBUSAN_LIST}}", build_tembusan_blocks())

    doc.save(out_path)
    print(f"OK: {out_path}")
    print(f"  paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
